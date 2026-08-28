from pathlib import Path

import docx
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


def iter_block_items(parent):
    """Yield Paragraph/Table children of parent in true document order.

    document.paragraphs and document.tables only expose flat, separate
    lists, so reading them one after another loses the real order when a
    table sits between paragraphs. This walks the underlying XML body (or
    a table cell) instead, so paragraphs and tables come out interleaved
    exactly as they appear in the document.
    """
    if isinstance(parent, DocxDocument):
        parent_element = parent.element.body
    elif isinstance(parent, _Cell):
        parent_element = parent._tc
    else:
        raise TypeError(
            "parent must be a Document or a table cell"
        )

    for child in parent_element.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def cell_text(cell: _Cell) -> str:
    return "\n".join(
        paragraph.text for paragraph in cell.paragraphs
    )


def extract_text_from_docx(file_path: Path) -> dict:
    document = docx.Document(file_path)

    lines = []

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            if block.text:
                lines.append(block.text)

        elif isinstance(block, Table):
            for row in block.rows:
                cell_texts = [
                    cell_text(cell)
                    for cell in row.cells
                    if cell_text(cell)
                ]

                if cell_texts:
                    lines.append(" | ".join(cell_texts))

    text = "\n".join(lines)

    return {
        "page_count": 1,
        "pages": [
            {
                "page_number": 1,
                "text": text,
            }
        ],
    }
