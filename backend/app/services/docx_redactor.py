from pathlib import Path

import docx

from backend.app.services.docx_locator import (
    build_docx_run_index,
    locate_entity_runs,
)
from backend.app.services.sensitive_value_masking import (
    mask_sensitive_value,
)


def _get_run(document, match: dict):
    if match["kind"] == "paragraph":
        paragraph = document.paragraphs[match["paragraph_index"]]
    else:
        table = document.tables[match["table_index"]]
        cell = table.rows[match["row_index"]].cells[
            match["cell_index"]
        ]
        paragraph = cell.paragraphs[match["paragraph_index"]]

    return paragraph.runs[match["run_index"]]


def _apply_mask_to_run_text(
    run_text: str,
    match_start: int,
    match_end: int,
    mask_slice: str,
) -> str:
    replacement_length = match_end - match_start

    if len(mask_slice) != replacement_length:
        mask_slice = "*" * replacement_length

    masked_characters = list(run_text)

    for offset, replacement in enumerate(mask_slice):
        run_offset = match_start + offset

        if (
            masked_characters[run_offset] == "*"
            or replacement == "*"
        ):
            masked_characters[run_offset] = "*"
        else:
            masked_characters[run_offset] = replacement

    return "".join(masked_characters)


def create_redacted_docx(
    source_path: Path,
    output_path: Path,
    findings: list[dict],
) -> Path:
    document = docx.Document(source_path)
    run_index = build_docx_run_index(source_path)

    for finding in findings:
        finding_start = finding["start"]
        finding_end = finding["end"]
        masked_value = mask_sensitive_value(
            finding["value"],
            finding["type"],
        )

        if len(masked_value) != finding_end - finding_start:
            masked_value = "*" * (finding_end - finding_start)

        matches = locate_entity_runs(
            run_index,
            finding_start,
            finding_end,
        )

        for match in matches:
            run = _get_run(document, match)
            overlap_start = (
                match["run_start"] + match["match_start"]
            )
            relative_start = overlap_start - finding_start
            relative_end = relative_start + (
                match["match_end"] - match["match_start"]
            )

            run.text = _apply_mask_to_run_text(
                run.text,
                match["match_start"],
                match["match_end"],
                masked_value[relative_start:relative_end],
            )

    document.save(output_path)

    return output_path
