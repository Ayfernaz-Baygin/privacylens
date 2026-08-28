import docx
import io

from fastapi.testclient import TestClient

from backend.app.main import app


def upload_docx(client, paragraphs):
    document = docx.Document()

    for paragraph in paragraphs:
        document.add_paragraph(paragraph)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = client.post(
        "/api/documents",
        files={
            "file": (
                "sample.docx",
                buffer,
                (
                    "application/vnd.openxmlformats"
                    "-officedocument.wordprocessingml"
                    ".document"
                ),
            )
        },
    )

    assert response.status_code == 200

    return response.json()["id"]


def test_analyze_docx_detects_findings_without_bounding_boxes(monkeypatch):
    monkeypatch.setattr(
        "backend.app.services.detection_engine.detect_named_entities",
        lambda text: [],
    )

    client = TestClient(app)

    document_id = upload_docx(
        client,
        ["E-mail: test@example.com"],
    )

    response = client.get(
        f"/api/documents/{document_id}/analyze"
    )

    assert response.status_code == 200

    body = response.json()

    assert body["status"] == "analyzed"
    assert body["page_count"] == 1
    assert body["finding_count"] == 1

    finding = body["findings"][0]

    assert finding["type"] == "EMAIL"
    assert finding["value"] == "test@example.com"
    assert finding["bounding_boxes"] == []
    assert finding["page_number"] == 1
    assert "finding_id" in finding
    assert "confidence_level" in finding
    assert "privacy_status" in finding
    assert "redaction_action" in finding


def test_analyze_docx_reads_table_cells(monkeypatch):
    monkeypatch.setattr(
        "backend.app.services.detection_engine.detect_named_entities",
        lambda text: [],
    )

    client = TestClient(app)

    document = docx.Document()
    document.add_paragraph("Müşteri bilgileri")

    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "test@example.com"

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    upload_response = client.post(
        "/api/documents",
        files={
            "file": (
                "sample-table.docx",
                buffer,
                (
                    "application/vnd.openxmlformats"
                    "-officedocument.wordprocessingml"
                    ".document"
                ),
            )
        },
    )

    document_id = upload_response.json()["id"]

    response = client.get(
        f"/api/documents/{document_id}/analyze"
    )

    assert response.status_code == 200

    body = response.json()

    assert body["finding_count"] == 1
    assert body["findings"][0]["type"] == "EMAIL"


def test_analyze_docx_without_findings_returns_empty_list(monkeypatch):
    monkeypatch.setattr(
        "backend.app.services.detection_engine.detect_named_entities",
        lambda text: [],
    )

    client = TestClient(app)

    document_id = upload_docx(
        client,
        ["Herhangi bir hassas veri içermeyen paragraf."],
    )

    response = client.get(
        f"/api/documents/{document_id}/analyze"
    )

    assert response.status_code == 200

    body = response.json()

    assert body["finding_count"] == 0
    assert body["findings"] == []
