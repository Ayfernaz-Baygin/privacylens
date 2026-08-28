import io
import zipfile

import docx
import pymupdf
from fastapi.testclient import TestClient

from backend.app.main import app


def upload_docx(client, text):
    document = docx.Document()
    document.add_paragraph(text)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = client.post(
        "/api/documents",
        files={
            "file": (
                "sample.docx",
                buffer,
                (
                    "application/vnd.openxmlformats"
                    "-officedocument.wordprocessingml"
                    ".document"
                ),
            )
        },
    )

    assert response.status_code == 200

    return response.json()["id"]


def fake_location_ner(text):
    start = text.index("İstanbul")
    end = start + len("İstanbul")

    return [
        {
            "type": "LOCATION",
            "value": "İstanbul",
            "start": start,
            "end": end,
            "confidence": 0.99,
            "source": "ner_model",
        }
    ]


def test_redact_selected_docx_auto_redacts_email(monkeypatch):
    monkeypatch.setattr(
        "backend.app.services.detection_engine.detect_named_entities",
        lambda text: [],
    )

    client = TestClient(app)

    document_id = upload_docx(
        client,
        "E-mail: test@example.com ile iletişime geçin.",
    )

    response = client.post(
        f"/api/documents/{document_id}/redact-selected",
        json={"selected_finding_ids": []},
    )

    assert response.status_code == 200

    redacted_document = docx.Document(io.BytesIO(response.content))
    redacted_text = redacted_document.paragraphs[0].text

    assert "test@example.com" not in redacted_text
    assert "ile iletişime geçin." in redacted_text
    assert "█" in redacted_text


def test_redact_selected_docx_keeps_unselected_review_finding(
    monkeypatch,
):
    monkeypatch.setattr(
        "backend.app.services.detection_engine.detect_named_entities",
        fake_location_ner,
    )

    client = TestClient(app)

    document_id = upload_docx(
        client,
        "Merkezimiz İstanbul şehrindedir. "
        "Email: test@example.com",
    )

    response = client.post(
        f"/api/documents/{document_id}/redact-selected",
        json={"selected_finding_ids": []},
    )

    assert response.status_code == 200

    redacted_document = docx.Document(io.BytesIO(response.content))
    redacted_text = redacted_document.paragraphs[0].text

    assert "İstanbul" in redacted_text
    assert "test@example.com" not in redacted_text


def test_redact_selected_docx_redacts_selected_review_finding(
    monkeypatch,
):
    monkeypatch.setattr(
        "backend.app.services.detection_engine.detect_named_entities",
        fake_location_ner,
    )

    client = TestClient(app)

    document_id = upload_docx(
        client,
        "Merkezimiz İstanbul şehrindedir. "
        "Email: test@example.com",
    )

    analyze_response = client.get(
        f"/api/documents/{document_id}/analyze"
    )

    assert analyze_response.status_code == 200

    findings = analyze_response.json()["findings"]

    location_finding = next(
        finding
        for finding in findings
        if finding["type"] == "LOCATION"
    )

    response = client.post(
        f"/api/documents/{document_id}/redact-selected",
        json={
            "selected_finding_ids": [
                location_finding["finding_id"]
            ]
        },
    )

    assert response.status_code == 200

    redacted_document = docx.Document(io.BytesIO(response.content))
    redacted_text = redacted_document.paragraphs[0].text

    assert "İstanbul" not in redacted_text
    assert "test@example.com" not in redacted_text


def test_redact_selected_docx_response_content_type(monkeypatch):
    monkeypatch.setattr(
        "backend.app.services.detection_engine.detect_named_entities",
        lambda text: [],
    )

    client = TestClient(app)

    document_id = upload_docx(
        client,
        "E-mail: test@example.com",
    )

    response = client.post(
        f"/api/documents/{document_id}/redact-selected",
        json={"selected_finding_ids": []},
    )

    assert response.status_code == 200
    assert response.headers["content-type"] == (
        "application/vnd.openxmlformats"
        "-officedocument.wordprocessingml.document"
    )
    assert response.headers["content-disposition"].endswith(
        '.docx"'
    )


def test_redact_selected_docx_removes_value_from_raw_xml(
    monkeypatch,
):
    monkeypatch.setattr(
        "backend.app.services.detection_engine.detect_named_entities",
        lambda text: [],
    )

    client = TestClient(app)

    document_id = upload_docx(
        client,
        "İletişim için ayfer.aycan@example.com "
        "adresini kullanabilirsiniz.",
    )

    response = client.post(
        f"/api/documents/{document_id}/redact-selected",
        json={"selected_finding_ids": []},
    )

    assert response.status_code == 200

    with zipfile.ZipFile(
        io.BytesIO(response.content)
    ) as archive:
        raw_xml = archive.read(
            "word/document.xml"
        ).decode("utf-8")

    assert "ayfer.aycan@example.com" not in raw_xml


def test_redact_selected_pdf_behavior_is_unchanged(monkeypatch):
    monkeypatch.setattr(
        "backend.app.services.detection_engine.detect_named_entities",
        lambda text: [],
    )

    client = TestClient(app)

    pdf_document = pymupdf.open()
    page = pdf_document.new_page()
    page.insert_text((72, 72), "E-mail: test@example.com")

    buffer = io.BytesIO()
    pdf_document.save(buffer)
    pdf_document.close()
    buffer.seek(0)

    upload_response = client.post(
        "/api/documents",
        files={
            "file": (
                "sample.pdf",
                buffer,
                "application/pdf",
            )
        },
    )

    assert upload_response.status_code == 200

    document_id = upload_response.json()["id"]

    response = client.post(
        f"/api/documents/{document_id}/redact-selected",
        json={"selected_finding_ids": []},
    )

    assert response.status_code == 200
    assert response.headers["content-type"] == "application/pdf"
    assert response.headers["content-disposition"].endswith(
        '.pdf"'
    )
