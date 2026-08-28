import docx

from backend.app.services.docx_parser import extract_text_from_docx


def test_extract_text_from_docx_reads_paragraphs(tmp_path):
    docx_path = tmp_path / "sample.docx"

    document = docx.Document()
    document.add_paragraph("Merhaba dünya")
    document.add_paragraph("Email: test@example.com")
    document.save(docx_path)

    result = extract_text_from_docx(docx_path)

    assert result["page_count"] == 1
    assert len(result["pages"]) == 1

    page = result["pages"][0]

    assert page["page_number"] == 1
    assert "Merhaba dünya" in page["text"]
    assert "Email: test@example.com" in page["text"]


def test_extract_text_from_docx_reads_table_cells(tmp_path):
    docx_path = tmp_path / "sample-table.docx"

    document = docx.Document()
    document.add_paragraph("Başlık")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Ad"
    table.cell(0, 1).text = "Soyad"
    table.cell(1, 0).text = "Ayfer"
    table.cell(1, 1).text = "Aycan"

    document.save(docx_path)

    result = extract_text_from_docx(docx_path)

    page = result["pages"][0]

    assert "Başlık" in page["text"]
    assert "Ad | Soyad" in page["text"]
    assert "Ayfer | Aycan" in page["text"]


def test_extract_text_from_docx_skips_empty_paragraphs(tmp_path):
    docx_path = tmp_path / "sample-empty.docx"

    document = docx.Document()
    document.add_paragraph("İlk paragraf")
    document.add_paragraph("")
    document.add_paragraph("İkinci paragraf")
    document.save(docx_path)

    result = extract_text_from_docx(docx_path)

    page = result["pages"][0]
    lines = page["text"].split("\n")

    assert lines == ["İlk paragraf", "İkinci paragraf"]
