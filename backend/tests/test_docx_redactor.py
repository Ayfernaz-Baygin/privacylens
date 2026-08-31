import zipfile

import docx

from backend.app.services.docx_parser import extract_text_from_docx
from backend.app.services.docx_redactor import create_redacted_docx


def _finding_for(
    text: str,
    value: str,
    finding_type: str,
) -> dict:
    start = text.index(value)

    return {
        "type": finding_type,
        "value": value,
        "start": start,
        "end": start + len(value),
    }


def test_create_redacted_docx_single_run(tmp_path):
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "redacted.docx"

    document = docx.Document()
    document.add_paragraph("E-mail: test@example.com")
    document.save(source_path)

    text = extract_text_from_docx(source_path)["pages"][0]["text"]
    value = "test@example.com"

    create_redacted_docx(
        source_path,
        output_path,
        [_finding_for(text, value, "EMAIL")],
    )

    redacted_document = docx.Document(output_path)
    paragraph = redacted_document.paragraphs[0]

    assert paragraph.text == "E-mail: t**t@e*****e.com"
    assert value not in paragraph.text
    assert len(paragraph.runs) == 1


def test_create_redacted_docx_entity_split_across_runs(tmp_path):
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "redacted.docx"

    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("E-posta: ayse@exa")
    paragraph.add_run("mple.com kayıtlı")
    document.save(source_path)

    text = extract_text_from_docx(source_path)["pages"][0]["text"]
    assert text == "E-posta: ayse@example.com kayıtlı"

    value = "ayse@example.com"

    create_redacted_docx(
        source_path,
        output_path,
        [_finding_for(text, value, "EMAIL")],
    )

    redacted_document = docx.Document(output_path)
    paragraph = redacted_document.paragraphs[0]

    assert len(paragraph.runs) == 2
    assert paragraph.runs[0].text == "E-posta: a**e@e**"
    assert paragraph.runs[1].text == "***e.com kayıtlı"
    assert paragraph.text == "E-posta: a**e@e*****e.com kayıtlı"
    assert value not in paragraph.text
    assert "kayıtlı" in paragraph.text


def test_create_redacted_docx_table_cell(tmp_path):
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "redacted.docx"

    document = docx.Document()
    document.add_paragraph("Müşteri bilgileri")

    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "E-posta"
    table.cell(0, 1).text = "test@example.com"
    document.save(source_path)

    text = extract_text_from_docx(source_path)["pages"][0]["text"]
    value = "test@example.com"

    create_redacted_docx(
        source_path,
        output_path,
        [_finding_for(text, value, "EMAIL")],
    )

    redacted_document = docx.Document(output_path)
    redacted_table = redacted_document.tables[0]

    assert redacted_table.cell(0, 1).text == "t**t@e*****e.com"
    assert redacted_table.cell(0, 0).text == "E-posta"
    assert redacted_document.paragraphs[0].text == "Müşteri bilgileri"


def test_create_redacted_docx_removes_sensitive_text_when_reopened(
    tmp_path,
):
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "redacted.docx"

    document = docx.Document()
    document.add_paragraph(
        "TCKN: 12345678901 ile başvuru yapıldı."
    )
    document.save(source_path)

    text = extract_text_from_docx(source_path)["pages"][0]["text"]
    value = "12345678901"

    create_redacted_docx(
        source_path,
        output_path,
        [_finding_for(text, value, "TCKN")],
    )

    redacted_text = extract_text_from_docx(output_path)[
        "pages"
    ][0]["text"]

    assert value not in redacted_text
    assert "***********" in redacted_text
    assert "TCKN:" in redacted_text
    assert "ile başvuru yapıldı." in redacted_text


def test_create_redacted_docx_preserves_neighboring_text(tmp_path):
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "redacted.docx"

    document = docx.Document()
    document.add_paragraph(
        "Ad Soyad: Ayfer Aycan, Email: test@example.com, "
        "Şehir: İstanbul"
    )
    document.save(source_path)

    text = extract_text_from_docx(source_path)["pages"][0]["text"]
    value = "test@example.com"

    create_redacted_docx(
        source_path,
        output_path,
        [_finding_for(text, value, "EMAIL")],
    )

    redacted_text = docx.Document(output_path).paragraphs[0].text

    assert "Ad Soyad: Ayfer Aycan" in redacted_text
    assert "Şehir: İstanbul" in redacted_text
    assert value not in redacted_text


def test_create_redacted_docx_preserves_run_formatting(tmp_path):
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "redacted.docx"

    document = docx.Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("test@example.com")
    run.bold = True
    run.italic = True
    document.save(source_path)

    text = extract_text_from_docx(source_path)["pages"][0]["text"]
    value = "test@example.com"

    create_redacted_docx(
        source_path,
        output_path,
        [_finding_for(text, value, "EMAIL")],
    )

    redacted_document = docx.Document(output_path)
    redacted_paragraph = redacted_document.paragraphs[0]

    assert len(redacted_paragraph.runs) == 1

    redacted_run = redacted_paragraph.runs[0]

    assert redacted_run.bold is True
    assert redacted_run.italic is True
    assert redacted_run.text == "t**t@e*****e.com"


def test_create_redacted_docx_masks_person_across_run_boundaries(
    tmp_path,
):
    source_path = tmp_path / "person-source.docx"
    output_path = tmp_path / "person-redacted.docx"

    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Kişi: Ay")
    paragraph.add_run("şe Yıl")
    paragraph.add_run("maz kayıtlı")
    document.save(source_path)

    text = extract_text_from_docx(source_path)["pages"][0]["text"]
    value = "Ayşe Yılmaz"

    create_redacted_docx(
        source_path,
        output_path,
        [_finding_for(text, value, "PERSON")],
    )

    redacted_paragraph = docx.Document(output_path).paragraphs[0]

    assert redacted_paragraph.text == "Kişi: A**e Y****z kayıtlı"
    assert [run.text for run in redacted_paragraph.runs] == [
        "Kişi: A*",
        "*e Y**",
        "**z kayıtlı",
    ]


def test_create_redacted_docx_masks_phone(tmp_path):
    source_path = tmp_path / "phone-source.docx"
    output_path = tmp_path / "phone-redacted.docx"

    document = docx.Document()
    document.add_paragraph("Telefon: 0532-123-45-67")
    document.save(source_path)

    text = extract_text_from_docx(source_path)["pages"][0]["text"]
    value = "0532-123-45-67"

    create_redacted_docx(
        source_path,
        output_path,
        [_finding_for(text, value, "PHONE")],
    )

    assert (
        docx.Document(output_path).paragraphs[0].text
        == "Telefon: 0***-***-**-*7"
    )


def test_overlapping_findings_never_reveal_a_more_restricted_mask(
    tmp_path,
):
    source_path = tmp_path / "overlap-source.docx"
    output_path = tmp_path / "overlap-redacted.docx"
    value = "12345678901"

    document = docx.Document()
    document.add_paragraph(value)
    document.save(source_path)

    text = extract_text_from_docx(source_path)["pages"][0]["text"]

    create_redacted_docx(
        source_path,
        output_path,
        [
            _finding_for(text, value, "TCKN"),
            _finding_for(text, value, "PHONE"),
        ],
    )

    assert docx.Document(output_path).paragraphs[0].text == "***********"


def test_create_redacted_docx_replaces_value_in_raw_xml(tmp_path):
    source_path = tmp_path / "xml-source.docx"
    output_path = tmp_path / "xml-redacted.docx"
    value = "ayfer.aycan@example.com"
    masked_value = "a*********n@e*****e.com"

    document = docx.Document()
    document.add_paragraph(f"İletişim: {value} adresini kullanın.")
    document.save(source_path)

    text = extract_text_from_docx(source_path)["pages"][0]["text"]
    create_redacted_docx(
        source_path,
        output_path,
        [_finding_for(text, value, "EMAIL")],
    )

    with zipfile.ZipFile(output_path) as archive:
        raw_xml = archive.read("word/document.xml").decode("utf-8")

    assert value not in raw_xml
    assert masked_value in raw_xml
    assert "adresini kullanın." in raw_xml
