import docx

from backend.app.services.docx_locator import (
    build_docx_run_index,
    locate_text_in_docx,
)
from backend.app.services.docx_parser import extract_text_from_docx


def test_locate_text_in_docx_single_run(tmp_path):
    docx_path = tmp_path / "single-run.docx"

    document = docx.Document()
    document.add_paragraph("E-mail: test@example.com")
    document.save(docx_path)

    text = extract_text_from_docx(docx_path)["pages"][0]["text"]
    start = text.index("test@example.com")
    end = start + len("test@example.com")

    matches = locate_text_in_docx(docx_path, start, end)

    assert len(matches) == 1

    match = matches[0]

    assert match["kind"] == "paragraph"
    assert match["paragraph_index"] == 0
    assert match["run_index"] == 0
    assert match["matched_text"] == "test@example.com"


def test_locate_text_in_docx_entity_split_across_runs(tmp_path):
    docx_path = tmp_path / "split-run.docx"

    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("E-mail: test@ex")
    paragraph.add_run("ample.com")
    document.save(docx_path)

    text = extract_text_from_docx(docx_path)["pages"][0]["text"]
    assert text == "E-mail: test@example.com"

    start = text.index("test@example.com")
    end = start + len("test@example.com")

    matches = locate_text_in_docx(docx_path, start, end)

    assert len(matches) == 2

    first, second = matches

    assert first["run_index"] == 0
    assert first["matched_text"] == "test@ex"

    assert second["run_index"] == 1
    assert second["matched_text"] == "ample.com"

    assert (
        first["matched_text"] + second["matched_text"]
        == "test@example.com"
    )


def test_locate_text_in_docx_table_cell(tmp_path):
    docx_path = tmp_path / "table-cell.docx"

    document = docx.Document()
    document.add_paragraph("Müşteri bilgileri")

    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "E-posta"
    table.cell(0, 1).text = "test@example.com"

    document.save(docx_path)

    text = extract_text_from_docx(docx_path)["pages"][0]["text"]
    start = text.index("test@example.com")
    end = start + len("test@example.com")

    matches = locate_text_in_docx(docx_path, start, end)

    assert len(matches) == 1

    match = matches[0]

    assert match["kind"] == "table_cell"
    assert match["table_index"] == 0
    assert match["row_index"] == 0
    assert match["cell_index"] == 1
    assert match["matched_text"] == "test@example.com"


def test_locate_text_in_docx_paragraph_table_paragraph_order(tmp_path):
    docx_path = tmp_path / "order.docx"

    document = docx.Document()
    document.add_paragraph("Before test@before.com")

    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "test@middle.com"

    document.add_paragraph("After test@after.com")
    document.save(docx_path)

    run_index = build_docx_run_index(docx_path)

    assert run_index["text"] == (
        "Before test@before.com\n"
        "test@middle.com\n"
        "After test@after.com"
    )

    for value, expected_kind in (
        ("test@before.com", "paragraph"),
        ("test@middle.com", "table_cell"),
        ("test@after.com", "paragraph"),
    ):
        start = run_index["text"].index(value)
        end = start + len(value)

        matches = locate_text_in_docx(docx_path, start, end)

        assert len(matches) == 1
        assert matches[0]["kind"] == expected_kind
        assert matches[0]["matched_text"] == value


def test_build_docx_run_index_matches_extract_text_from_docx(tmp_path):
    docx_path = tmp_path / "parity.docx"

    document = docx.Document()
    document.add_paragraph("Başlık")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Ad"
    table.cell(0, 1).text = "Soyad"
    table.cell(1, 0).text = "Ayfer"
    table.cell(1, 1).text = "Aycan"

    document.add_paragraph("Son paragraf")
    document.save(docx_path)

    parsed_text = extract_text_from_docx(docx_path)["pages"][0]["text"]
    run_index = build_docx_run_index(docx_path)

    assert run_index["text"] == parsed_text
