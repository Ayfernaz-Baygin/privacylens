import io
import os
import time
import uuid
from pathlib import Path

import docx
import openpyxl
import pymupdf
import pytest
from fastapi.testclient import TestClient

from backend.app import main as main_module
from backend.app.main import app
from backend.app.routes.documents import UPLOAD_ROOT
from backend.app.services.document_cleanup import (
    cleanup_stale_documents,
    delete_document_directory,
)

PDF_CONTENT_TYPE = "application/pdf"
DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats"
    "-officedocument.wordprocessingml.document"
)
XLSX_CONTENT_TYPE = (
    "application/vnd.openxmlformats"
    "-officedocument.spreadsheetml.sheet"
)


def _no_ner(monkeypatch):
    monkeypatch.setattr(
        "backend.app.services.detection_engine.detect_named_entities",
        lambda text: [],
    )


def _real_pdf_bytes() -> bytes:
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "E-mail: test@example.com")

    buffer = io.BytesIO()
    document.save(buffer)
    document.close()

    return buffer.getvalue()


def _real_docx_bytes() -> bytes:
    document = docx.Document()
    document.add_paragraph("E-mail: test@example.com")

    buffer = io.BytesIO()
    document.save(buffer)

    return buffer.getvalue()


def _real_xlsx_bytes() -> bytes:
    workbook = openpyxl.Workbook()
    workbook.active["A1"] = "E-mail: test@example.com"

    buffer = io.BytesIO()
    workbook.save(buffer)

    return buffer.getvalue()


def _upload(client, filename, content, content_type) -> str:
    response = client.post(
        "/api/documents",
        files={
            "file": (filename, io.BytesIO(content), content_type)
        },
    )

    assert response.status_code == 200

    return response.json()["id"]


def _touch_old(path: Path, age_seconds: float) -> None:
    old_time = time.time() - age_seconds
    os.utime(path, (old_time, old_time))


# --- delete_document_directory -----------------------------------------


def test_delete_document_directory_removes_target(tmp_path):
    document_id = str(uuid.uuid4())
    target = tmp_path / document_id
    target.mkdir()
    (target / "source.pdf").write_bytes(b"data")

    delete_document_directory(tmp_path, document_id)

    assert not target.exists()


def test_delete_document_directory_does_not_touch_other_directories(
    tmp_path,
):
    document_id = str(uuid.uuid4())
    other_id = str(uuid.uuid4())

    target = tmp_path / document_id
    target.mkdir()

    other = tmp_path / other_id
    other.mkdir()
    (other / "source.pdf").write_bytes(b"data")

    delete_document_directory(tmp_path, document_id)

    assert not target.exists()
    assert other.exists()
    assert (other / "source.pdf").exists()


def test_delete_document_directory_is_a_no_op_for_missing_directory(
    tmp_path,
):
    document_id = str(uuid.uuid4())

    delete_document_directory(tmp_path, document_id)


# --- DELETE /api/documents/{document_id} --------------------------------


def test_delete_endpoint_removes_valid_document():
    client = TestClient(app)

    document_id = _upload(
        client, "sample.pdf", _real_pdf_bytes(), PDF_CONTENT_TYPE
    )

    response = client.delete(
        f"/api/documents/{document_id}"
    )

    assert response.status_code == 204
    assert not (UPLOAD_ROOT / document_id).exists()


def test_delete_endpoint_returns_404_for_unknown_valid_uuid():
    client = TestClient(app)

    unknown_id = str(uuid.uuid4())

    response = client.delete(
        f"/api/documents/{unknown_id}"
    )

    assert response.status_code == 404


def test_delete_endpoint_rejects_malformed_document_id():
    client = TestClient(app)

    response = client.delete(
        "/api/documents/not-a-uuid"
    )

    assert response.status_code == 404


def test_delete_endpoint_rejects_percent_encoded_traversal_id():
    client = TestClient(app)

    response = client.delete(
        "/api/documents/..%2F..%2Fetc"
    )

    assert response.status_code == 404


# --- /redact-selected cleanup, per format --------------------------------


def test_redact_selected_pdf_cleans_up_directory_after_response(
    monkeypatch,
):
    _no_ner(monkeypatch)

    client = TestClient(app)

    document_id = _upload(
        client, "sample.pdf", _real_pdf_bytes(), PDF_CONTENT_TYPE
    )

    assert (UPLOAD_ROOT / document_id).exists()

    response = client.post(
        f"/api/documents/{document_id}/redact-selected",
        json={"selected_finding_ids": []},
    )

    assert response.status_code == 200
    assert response.headers["content-type"] == "application/pdf"
    assert len(response.content) > 0

    assert not (UPLOAD_ROOT / document_id).exists()


def test_redact_selected_docx_cleans_up_directory_after_response(
    monkeypatch,
):
    _no_ner(monkeypatch)

    client = TestClient(app)

    document_id = _upload(
        client, "sample.docx", _real_docx_bytes(), DOCX_CONTENT_TYPE
    )

    assert (UPLOAD_ROOT / document_id).exists()

    response = client.post(
        f"/api/documents/{document_id}/redact-selected",
        json={"selected_finding_ids": []},
    )

    assert response.status_code == 200

    redacted_document = docx.Document(io.BytesIO(response.content))
    assert "test@example.com" not in redacted_document.paragraphs[0].text

    assert not (UPLOAD_ROOT / document_id).exists()


def test_redact_selected_xlsx_cleans_up_directory_after_response(
    monkeypatch,
):
    _no_ner(monkeypatch)

    client = TestClient(app)

    document_id = _upload(
        client, "sample.xlsx", _real_xlsx_bytes(), XLSX_CONTENT_TYPE
    )

    assert (UPLOAD_ROOT / document_id).exists()

    response = client.post(
        f"/api/documents/{document_id}/redact-selected",
        json={"selected_finding_ids": []},
    )

    assert response.status_code == 200

    redacted = openpyxl.load_workbook(io.BytesIO(response.content))
    assert "test@example.com" not in redacted.active["A1"].value

    assert not (UPLOAD_ROOT / document_id).exists()


# --- cleanup_stale_documents (TTL sweep) ---------------------------------


def test_cleanup_stale_documents_removes_directory_older_than_retention(
    tmp_path,
):
    document_id = str(uuid.uuid4())
    target = tmp_path / document_id
    target.mkdir()
    source = target / "source.pdf"
    source.write_bytes(b"data")

    _touch_old(source, age_seconds=7200)
    _touch_old(target, age_seconds=7200)

    deleted_count = cleanup_stale_documents(
        tmp_path, retention_seconds=3600
    )

    assert deleted_count == 1
    assert not target.exists()


def test_cleanup_stale_documents_keeps_fresh_directory(tmp_path):
    document_id = str(uuid.uuid4())
    target = tmp_path / document_id
    target.mkdir()
    (target / "source.pdf").write_bytes(b"data")

    deleted_count = cleanup_stale_documents(
        tmp_path, retention_seconds=3600
    )

    assert deleted_count == 0
    assert target.exists()


def test_cleanup_stale_documents_ignores_non_uuid_directory(tmp_path):
    unexpected = tmp_path / "not-a-document-id"
    unexpected.mkdir()
    stray_file = unexpected / "whatever.txt"
    stray_file.write_bytes(b"data")

    _touch_old(stray_file, age_seconds=7200)
    _touch_old(unexpected, age_seconds=7200)

    deleted_count = cleanup_stale_documents(
        tmp_path, retention_seconds=3600
    )

    assert deleted_count == 0
    assert unexpected.exists()


def test_cleanup_stale_documents_one_failure_does_not_block_others(
    tmp_path, monkeypatch,
):
    failing_id = str(uuid.uuid4())
    ok_id = str(uuid.uuid4())

    failing_dir = tmp_path / failing_id
    failing_dir.mkdir()
    ok_dir = tmp_path / ok_id
    ok_dir.mkdir()

    for directory in (failing_dir, ok_dir):
        source = directory / "source.pdf"
        source.write_bytes(b"data")
        _touch_old(source, age_seconds=7200)
        _touch_old(directory, age_seconds=7200)

    import shutil

    real_rmtree = shutil.rmtree

    def fake_rmtree(path, *args, **kwargs):
        if Path(path) == failing_dir:
            raise OSError("simulated failure")
        return real_rmtree(path, *args, **kwargs)

    monkeypatch.setattr(
        "backend.app.services.document_cleanup.shutil.rmtree",
        fake_rmtree,
    )

    deleted_count = cleanup_stale_documents(
        tmp_path, retention_seconds=3600
    )

    assert deleted_count == 1
    assert failing_dir.exists()
    assert not ok_dir.exists()


# --- periodic background cleanup task lifecycle --------------------------


def test_periodic_cleanup_task_removes_stale_directory(monkeypatch):
    monkeypatch.setattr(
        main_module, "CLEANUP_INTERVAL_SECONDS", 0.05
    )
    monkeypatch.setattr(
        main_module, "DOCUMENT_RETENTION_SECONDS", 0
    )

    with TestClient(app):
        # Created only after the app is already running (past the
        # one-shot startup sweep), so only the periodic loop -- not
        # lifespan startup -- can be what removes it.
        document_id = str(uuid.uuid4())
        target = UPLOAD_ROOT / document_id
        target.mkdir(parents=True, exist_ok=True)
        (target / "source.pdf").write_bytes(b"data")

        deadline = time.time() + 2.0

        while time.time() < deadline and target.exists():
            time.sleep(0.05)

        assert not target.exists()


def test_periodic_cleanup_task_is_cancelled_on_shutdown(monkeypatch):
    monkeypatch.setattr(
        main_module, "CLEANUP_INTERVAL_SECONDS", 0.05
    )

    with TestClient(app):
        pass

    # Shutdown must complete (lifespan's `await cleanup_task` after
    # cancel doesn't hang) -- reaching this line is the assertion.
