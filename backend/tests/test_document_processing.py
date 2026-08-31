import docx
import openpyxl
import pymupdf
import pytest

from backend.app.services.document_processing import (
    DocumentParseError,
    UnsupportedDocumentFormatError,
    analyze_document_file,
    build_finding_id,
    locate_finding_in_hybrid_regions,
    locate_finding_in_ocr_regions,
)
from backend.app.services.pdf_parser import PdfOcrError


def _no_ner(monkeypatch):
    monkeypatch.setattr(
        "backend.app.services.detection_engine.detect_named_entities",
        lambda text: [],
    )


def test_build_finding_id_combines_page_start_end_type():
    finding = {
        "page_number": 3,
        "start": 10,
        "end": 20,
        "type": "EMAIL",
    }

    assert build_finding_id(finding) == "3:10:20:EMAIL"


def _mock_pdf_analysis(
    tmp_path,
    monkeypatch,
    pages,
    findings_by_text,
):
    pdf_path = tmp_path / "source.pdf"
    pdf_path.write_bytes(b"mock pdf")

    monkeypatch.setattr(
        "backend.app.services.document_processing.extract_text_from_pdf",
        lambda file_path: {
            "page_count": len(pages),
            "pages": pages,
        },
    )
    monkeypatch.setattr(
        "backend.app.services.document_processing.detect_sensitive_data",
        lambda text, page_number, include_ner: [
            {
                **finding,
                "page_number": page_number,
            }
            for finding in findings_by_text.get(text, [])
        ],
    )

    return analyze_document_file(
        pdf_path=pdf_path,
        docx_path=tmp_path / "source.docx",
        xlsx_path=tmp_path / "source.xlsx",
    )


def test_ocr_finding_gets_intersecting_region_bbox(
    tmp_path, monkeypatch
):
    result = _mock_pdf_analysis(
        tmp_path,
        monkeypatch,
        pages=[
            {
                "page_number": 1,
                "text": "Email test@example.com\n",
                "text_source": "ocr",
                "regions": [
                    {
                        "start": 6,
                        "end": 22,
                        "text": "test@example.com",
                        "bbox": {
                            "x0": 10,
                            "y0": 20,
                            "x1": 100,
                            "y1": 32,
                        },
                    }
                ],
            }
        ],
        findings_by_text={
            "Email test@example.com\n": [
                {
                    "type": "EMAIL",
                    "value": "test@example.com",
                    "start": 6,
                    "end": 22,
                }
            ]
        },
    )

    assert result["findings"][0]["bounding_boxes"] == [
        {"x0": 10, "y0": 20, "x1": 100, "y1": 32}
    ]


def test_repeated_ocr_text_only_matches_finding_occurrence():
    first_box = {"x0": 1, "y0": 2, "x1": 3, "y1": 4}
    second_box = {"x0": 5, "y0": 6, "x1": 7, "y1": 8}

    boxes = locate_finding_in_ocr_regions(
        finding={"start": 23, "end": 39},
        regions=[
            {
                "start": 0,
                "end": 16,
                "text": "test@example.com",
                "bbox": first_box,
            },
            {
                "start": 23,
                "end": 39,
                "text": "test@example.com",
                "bbox": second_box,
            },
        ],
    )

    assert boxes == [second_box]


def test_ocr_finding_can_span_multiple_regions():
    boxes = locate_finding_in_ocr_regions(
        finding={"start": 3, "end": 12},
        regions=[
            {
                "start": 0,
                "end": 5,
                "text": "Ahmet",
                "bbox": {"x0": 1, "y0": 2, "x1": 3, "y1": 4},
            },
            {
                "start": 5,
                "end": 12,
                "text": " Yilmaz",
                "bbox": {"x0": 5, "y0": 6, "x1": 7, "y1": 8},
            },
        ],
    )

    assert boxes == [
        {"x0": 1, "y0": 2, "x1": 3, "y1": 4},
        {"x0": 5, "y0": 6, "x1": 7, "y1": 8},
    ]


def test_hybrid_native_region_finding_gets_correct_bbox(
    tmp_path, monkeypatch
):
    native_box = {"x0": 1, "y0": 2, "x1": 3, "y1": 4}
    monkeypatch.setattr(
        "backend.app.services.document_processing.locate_text_in_pdf",
        lambda **kwargs: pytest.fail(
            "hybrid page must not use the native PDF locator"
        ),
    )

    result = _mock_pdf_analysis(
        tmp_path,
        monkeypatch,
        pages=[
            {
                "page_number": 1,
                "text": "native@example.com\n|\nocr@example.com",
                "text_source": "hybrid",
                "regions": [
                    {
                        "start": 0,
                        "end": 18,
                        "text": "native@example.com",
                        "source": "native",
                        "bbox": native_box,
                    },
                    {
                        "start": 21,
                        "end": 36,
                        "text": "ocr@example.com",
                        "source": "ocr",
                        "bbox": {
                            "x0": 10,
                            "y0": 20,
                            "x1": 30,
                            "y1": 40,
                        },
                    },
                ],
            }
        ],
        findings_by_text={
            "native@example.com\n|\nocr@example.com": [
                {
                    "type": "EMAIL",
                    "value": "native@example.com",
                    "start": 0,
                    "end": 18,
                }
            ]
        },
    )

    assert result["findings"][0]["bounding_boxes"] == [native_box]


def test_hybrid_ocr_region_finding_gets_correct_bbox(
    tmp_path, monkeypatch
):
    ocr_box = {"x0": 10, "y0": 20, "x1": 30, "y1": 40}
    monkeypatch.setattr(
        "backend.app.services.document_processing.locate_text_in_pdf",
        lambda **kwargs: pytest.fail(
            "hybrid page must not use the native PDF locator"
        ),
    )

    result = _mock_pdf_analysis(
        tmp_path,
        monkeypatch,
        pages=[
            {
                "page_number": 1,
                "text": "Header\n|\nocr@example.com",
                "text_source": "hybrid",
                "regions": [
                    {
                        "start": 0,
                        "end": 6,
                        "text": "Header",
                        "source": "native",
                        "bbox": {
                            "x0": 1,
                            "y0": 2,
                            "x1": 3,
                            "y1": 4,
                        },
                    },
                    {
                        "start": 9,
                        "end": 24,
                        "text": "ocr@example.com",
                        "source": "ocr",
                        "bbox": ocr_box,
                    },
                ],
            }
        ],
        findings_by_text={
            "Header\n|\nocr@example.com": [
                {
                    "type": "EMAIL",
                    "value": "ocr@example.com",
                    "start": 9,
                    "end": 24,
                }
            ]
        },
    )

    assert result["findings"][0]["bounding_boxes"] == [ocr_box]


def test_hybrid_finding_can_span_native_and_ocr_regions():
    boxes = locate_finding_in_hybrid_regions(
        finding={"start": 3, "end": 12},
        regions=[
            {
                "start": 0,
                "end": 5,
                "source": "native",
                "bbox": {"x0": 1, "y0": 2, "x1": 3, "y1": 4},
            },
            {
                "start": 5,
                "end": 12,
                "source": "ocr",
                "bbox": {"x0": 5, "y0": 6, "x1": 7, "y1": 8},
            },
        ],
    )

    assert boxes == [
        {"x0": 1, "y0": 2, "x1": 3, "y1": 4},
        {"x0": 5, "y0": 6, "x1": 7, "y1": 8},
    ]


def test_hybrid_mapping_deduplicates_identical_bboxes():
    duplicate_box = {"x0": 1, "y0": 2, "x1": 3, "y1": 4}

    boxes = locate_finding_in_hybrid_regions(
        finding={"start": 0, "end": 10},
        regions=[
            {
                "start": 0,
                "end": 5,
                "source": "native",
                "bbox": duplicate_box,
            },
            {
                "start": 5,
                "end": 10,
                "source": "ocr",
                "bbox": duplicate_box,
            },
        ],
    )

    assert boxes == [duplicate_box]


def test_hybrid_repeated_text_only_maps_finding_occurrence():
    first_box = {"x0": 1, "y0": 2, "x1": 3, "y1": 4}
    second_box = {"x0": 5, "y0": 6, "x1": 7, "y1": 8}

    boxes = locate_finding_in_hybrid_regions(
        finding={"start": 20, "end": 36},
        regions=[
            {
                "start": 0,
                "end": 16,
                "text": "test@example.com",
                "source": "native",
                "bbox": first_box,
            },
            {
                "start": 20,
                "end": 36,
                "text": "test@example.com",
                "source": "ocr",
                "bbox": second_box,
            },
        ],
    )

    assert boxes == [second_box]


def test_native_pdf_still_uses_pdf_locator(tmp_path, monkeypatch):
    calls = []
    expected_box = {"x0": 1, "y0": 2, "x1": 3, "y1": 4}
    monkeypatch.setattr(
        "backend.app.services.document_processing.locate_text_in_pdf",
        lambda **kwargs: calls.append(kwargs) or [expected_box],
    )

    result = _mock_pdf_analysis(
        tmp_path,
        monkeypatch,
        pages=[
            {
                "page_number": 1,
                "text": "test@example.com",
                "text_source": "native",
                "regions": [],
            }
        ],
        findings_by_text={
            "test@example.com": [
                {
                    "type": "EMAIL",
                    "value": "test@example.com",
                    "start": 0,
                    "end": 16,
                }
            ]
        },
    )

    assert result["findings"][0]["bounding_boxes"] == [expected_box]
    assert calls[0]["page_number"] == 1
    assert calls[0]["value"] == "test@example.com"


def test_mixed_pdf_uses_native_ocr_and_hybrid_mapping(
    tmp_path, monkeypatch
):
    native_box = {"x0": 1, "y0": 2, "x1": 3, "y1": 4}
    ocr_box = {"x0": 10, "y0": 20, "x1": 30, "y1": 40}
    hybrid_box = {"x0": 50, "y0": 60, "x1": 70, "y1": 80}
    locator_pages = []
    monkeypatch.setattr(
        "backend.app.services.document_processing.locate_text_in_pdf",
        lambda **kwargs: locator_pages.append(
            kwargs["page_number"]
        ) or [native_box],
    )

    result = _mock_pdf_analysis(
        tmp_path,
        monkeypatch,
        pages=[
            {
                "page_number": 1,
                "text": "native@example.com",
                "text_source": "native",
                "regions": [],
            },
            {
                "page_number": 2,
                "text": "ocr@example.com",
                "text_source": "ocr",
                "regions": [
                    {
                        "start": 0,
                        "end": 15,
                        "text": "ocr@example.com",
                        "bbox": ocr_box,
                    }
                ],
            },
            {
                "page_number": 3,
                "text": "hybrid@example.com",
                "text_source": "hybrid",
                "regions": [
                    {
                        "start": 0,
                        "end": 18,
                        "text": "hybrid@example.com",
                        "source": "ocr",
                        "bbox": hybrid_box,
                    }
                ],
            },
        ],
        findings_by_text={
            "native@example.com": [
                {
                    "type": "EMAIL",
                    "value": "native@example.com",
                    "start": 0,
                    "end": 18,
                }
            ],
            "ocr@example.com": [
                {
                    "type": "EMAIL",
                    "value": "ocr@example.com",
                    "start": 0,
                    "end": 15,
                }
            ],
            "hybrid@example.com": [
                {
                    "type": "EMAIL",
                    "value": "hybrid@example.com",
                    "start": 0,
                    "end": 18,
                }
            ],
        },
    )

    assert locator_pages == [1]
    assert result["findings"][0]["bounding_boxes"] == [native_box]
    assert result["findings"][1]["bounding_boxes"] == [ocr_box]
    assert result["findings"][2]["bounding_boxes"] == [hybrid_box]


def test_analyze_document_file_uses_pdf_and_locates_boxes(
    tmp_path, monkeypatch
):
    _no_ner(monkeypatch)

    pdf_path = tmp_path / "source.pdf"
    docx_path = tmp_path / "source.docx"
    xlsx_path = tmp_path / "source.xlsx"

    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "E-mail: test@example.com")
    document.save(pdf_path)
    document.close()

    result = analyze_document_file(
        pdf_path=pdf_path,
        docx_path=docx_path,
        xlsx_path=xlsx_path,
    )

    assert result["document_format"] == "pdf"
    assert result["page_count"] == 1
    assert len(result["findings"]) == 1

    finding = result["findings"][0]

    assert finding["type"] == "EMAIL"
    assert len(finding["bounding_boxes"]) == 1
    assert finding["finding_id"] == build_finding_id(finding)


def test_analyze_document_file_uses_docx_with_empty_boxes(
    tmp_path, monkeypatch
):
    _no_ner(monkeypatch)

    pdf_path = tmp_path / "source.pdf"
    docx_path = tmp_path / "source.docx"
    xlsx_path = tmp_path / "source.xlsx"

    document = docx.Document()
    document.add_paragraph("E-mail: test@example.com")
    document.save(docx_path)

    result = analyze_document_file(
        pdf_path=pdf_path,
        docx_path=docx_path,
        xlsx_path=xlsx_path,
    )

    assert result["document_format"] == "docx"
    assert result["page_count"] == 1
    assert len(result["findings"]) == 1

    finding = result["findings"][0]

    assert finding["type"] == "EMAIL"
    assert finding["bounding_boxes"] == []
    assert finding["finding_id"] == build_finding_id(finding)


def test_analyze_document_file_uses_xlsx_with_empty_boxes(
    tmp_path, monkeypatch
):
    _no_ner(monkeypatch)

    pdf_path = tmp_path / "source.pdf"
    docx_path = tmp_path / "source.docx"
    xlsx_path = tmp_path / "source.xlsx"

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Müşteriler"
    sheet["A1"] = "E-mail: test@example.com"
    workbook.save(xlsx_path)

    result = analyze_document_file(
        pdf_path=pdf_path,
        docx_path=docx_path,
        xlsx_path=xlsx_path,
    )

    assert result["document_format"] == "xlsx"
    assert result["page_count"] == 1
    assert len(result["findings"]) == 1

    finding = result["findings"][0]

    assert finding["type"] == "EMAIL"
    assert finding["bounding_boxes"] == []
    assert finding["page_number"] == 1
    assert finding["finding_id"] == build_finding_id(finding)


def test_analyze_document_file_prefers_pdf_when_all_exist(
    tmp_path, monkeypatch
):
    _no_ner(monkeypatch)

    pdf_path = tmp_path / "source.pdf"
    docx_path = tmp_path / "source.docx"
    xlsx_path = tmp_path / "source.xlsx"

    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "test@example.com")
    document.save(pdf_path)
    document.close()

    docx.Document().save(docx_path)
    openpyxl.Workbook().save(xlsx_path)

    result = analyze_document_file(
        pdf_path=pdf_path,
        docx_path=docx_path,
        xlsx_path=xlsx_path,
    )

    assert result["document_format"] == "pdf"


def test_analyze_document_file_prefers_docx_over_xlsx(
    tmp_path, monkeypatch
):
    _no_ner(monkeypatch)

    pdf_path = tmp_path / "source.pdf"
    docx_path = tmp_path / "source.docx"
    xlsx_path = tmp_path / "source.xlsx"

    docx.Document().save(docx_path)
    openpyxl.Workbook().save(xlsx_path)

    result = analyze_document_file(
        pdf_path=pdf_path,
        docx_path=docx_path,
        xlsx_path=xlsx_path,
    )

    assert result["document_format"] == "docx"


def test_analyze_document_file_raises_when_none_exist(tmp_path):
    with pytest.raises(UnsupportedDocumentFormatError):
        analyze_document_file(
            pdf_path=tmp_path / "source.pdf",
            docx_path=tmp_path / "source.docx",
            xlsx_path=tmp_path / "source.xlsx",
        )


def test_analyze_document_file_raises_parse_error_for_corrupt_pdf(
    tmp_path,
):
    pdf_path = tmp_path / "source.pdf"
    docx_path = tmp_path / "source.docx"
    xlsx_path = tmp_path / "source.xlsx"

    pdf_path.write_bytes(b"not a real pdf")

    with pytest.raises(DocumentParseError) as excinfo:
        analyze_document_file(
            pdf_path=pdf_path,
            docx_path=docx_path,
            xlsx_path=xlsx_path,
        )

    assert excinfo.value.document_format == "pdf"


def test_analyze_document_file_preserves_ocr_error_detail(
    tmp_path, monkeypatch
):
    pdf_path = tmp_path / "source.pdf"
    pdf_path.write_bytes(b"mock pdf")

    def raise_ocr_error(file_path):
        raise PdfOcrError(page_number=3)

    monkeypatch.setattr(
        "backend.app.services.document_processing.extract_text_from_pdf",
        raise_ocr_error,
    )

    with pytest.raises(DocumentParseError) as excinfo:
        analyze_document_file(
            pdf_path=pdf_path,
            docx_path=tmp_path / "source.docx",
            xlsx_path=tmp_path / "source.xlsx",
        )

    assert excinfo.value.document_format == "pdf"
    assert excinfo.value.detail is not None
    assert "PDF page 3" in excinfo.value.detail
    assert "Tesseract" in excinfo.value.detail
    assert isinstance(excinfo.value.__cause__, PdfOcrError)


def test_analyze_document_file_raises_parse_error_for_corrupt_docx(
    tmp_path,
):
    pdf_path = tmp_path / "source.pdf"
    docx_path = tmp_path / "source.docx"
    xlsx_path = tmp_path / "source.xlsx"

    docx_path.write_bytes(b"not a real docx")

    with pytest.raises(DocumentParseError) as excinfo:
        analyze_document_file(
            pdf_path=pdf_path,
            docx_path=docx_path,
            xlsx_path=xlsx_path,
        )

    assert excinfo.value.document_format == "docx"


def test_analyze_document_file_raises_parse_error_for_corrupt_xlsx(
    tmp_path,
):
    pdf_path = tmp_path / "source.pdf"
    docx_path = tmp_path / "source.docx"
    xlsx_path = tmp_path / "source.xlsx"

    xlsx_path.write_bytes(b"not a real xlsx")

    with pytest.raises(DocumentParseError) as excinfo:
        analyze_document_file(
            pdf_path=pdf_path,
            docx_path=docx_path,
            xlsx_path=xlsx_path,
        )

    assert excinfo.value.document_format == "xlsx"
